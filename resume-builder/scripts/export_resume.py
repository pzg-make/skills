#!/usr/bin/env python3
"""Generate editable resume outputs from structured data or Markdown.

This script is intended for new or rebuilt resumes. For imported DOCX
preservation edits, generate a new document from the source structure instead
of modifying or overwriting the imported file.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
from pathlib import Path
from typing import Any


def load_input(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    suffix = path.suffix.lower()
    if suffix == ".json":
        data = json.loads(text)
    elif suffix in {".yaml", ".yml"}:
        data = load_yaml(text)
    elif suffix in {".md", ".markdown", ".txt"}:
        data = parse_markdown(text)
    else:
        raise SystemExit(f"Unsupported input type: {path.suffix}")
    if not isinstance(data, dict):
        raise SystemExit("Resume input must resolve to an object/dictionary.")
    return normalize_resume(data)


def load_yaml(text: str) -> Any:
    try:
        import yaml  # type: ignore

        return yaml.safe_load(text)
    except ModuleNotFoundError:
        return parse_simple_yaml(text)


def parse_simple_yaml(text: str) -> dict[str, Any]:
    """Parse a small YAML subset: top-level scalars and hyphen lists.

    This fallback keeps the script useful without PyYAML. Complex YAML should
    be converted to JSON or used with an environment that has PyYAML installed.
    """
    result: dict[str, Any] = {}
    current_key: str | None = None
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        if not line.startswith(" ") and ":" in line:
            key, value = line.split(":", 1)
            current_key = key.strip()
            value = value.strip()
            if value:
                result[current_key] = coerce_scalar(value)
            else:
                result[current_key] = []
            continue
        if current_key and line.strip().startswith("- "):
            item = line.strip()[2:].strip()
            if not isinstance(result.get(current_key), list):
                result[current_key] = []
            result[current_key].append(coerce_scalar(item))
    return result


def coerce_scalar(value: str) -> Any:
    value = value.strip().strip('"').strip("'")
    if value.lower() in {"true", "false"}:
        return value.lower() == "true"
    return value


def parse_markdown(text: str) -> dict[str, Any]:
    lines = text.splitlines()
    name = ""
    headline = ""
    contact: list[str] = []
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            name = stripped[2:].strip()
        elif stripped.startswith("## "):
            current = {"title": stripped[3:].strip(), "items": []}
            sections.append(current)
        elif current is None and not headline:
            headline = stripped
        elif current is None:
            contact.extend(split_contact(stripped))
        elif stripped.startswith(("- ", "* ", "+ ")):
            current["items"].append(stripped[2:].strip())
        else:
            current["items"].append(stripped)
    return {"name": name, "headline": headline, "contact": contact, "sections": sections}


def split_contact(value: str) -> list[str]:
    return [part.strip() for part in re.split(r"\s+[|/]\s+|,", value) if part.strip()]


def normalize_resume(data: dict[str, Any]) -> dict[str, Any]:
    contact = data.get("contact", [])
    if isinstance(contact, dict):
        contact = [str(v) for v in contact.values() if v]
    elif isinstance(contact, str):
        contact = split_contact(contact)

    sections = list(data.get("sections") or [])
    known_sections = [
        ("Summary", data.get("summary")),
        ("Skills", data.get("skills")),
        ("Experience", data.get("experience")),
        ("Projects", data.get("projects")),
        ("Education", data.get("education")),
        ("Certifications", data.get("certifications")),
        ("Awards", data.get("awards")),
    ]
    existing_titles = {str(section.get("title", "")).lower() for section in sections if isinstance(section, dict)}
    for title, value in known_sections:
        if value and title.lower() not in existing_titles:
            sections.append({"title": title, "items": ensure_items(value)})

    return {
        "name": str(data.get("name") or data.get("full_name") or "Resume").strip(),
        "headline": str(data.get("headline") or data.get("title") or "").strip(),
        "contact": [str(item).strip() for item in ensure_items(contact) if str(item).strip()],
        "sections": [normalize_section(section) for section in sections if section],
    }


def normalize_section(section: Any) -> dict[str, Any]:
    if isinstance(section, str):
        return {"title": section, "items": []}
    if not isinstance(section, dict):
        return {"title": "Section", "items": ensure_items(section)}
    title = str(section.get("title") or section.get("name") or "Section").strip()
    items = section.get("items", section.get("bullets", section.get("content", [])))
    return {"title": title, "items": ensure_items(items)}


def ensure_items(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    if isinstance(value, dict):
        return [value]
    return [value]


def item_to_lines(item: Any) -> tuple[str | None, list[str]]:
    if isinstance(item, dict):
        title_parts = [
            item.get("role") or item.get("title") or item.get("name"),
            item.get("company") or item.get("organization"),
            item.get("date") or item.get("dates") or item.get("period"),
        ]
        title = " | ".join(str(part).strip() for part in title_parts if part)
        bullets = item.get("bullets", item.get("items", item.get("description", [])))
        return title or None, [str(value).strip() for value in ensure_items(bullets) if str(value).strip()]
    return None, [str(item).strip()]


def render_markdown(resume: dict[str, Any]) -> str:
    out: list[str] = [f"# {resume['name']}"]
    if resume["headline"]:
        out.append(resume["headline"])
    if resume["contact"]:
        out.append(" | ".join(resume["contact"]))
    for section in resume["sections"]:
        out.extend(["", f"## {section['title']}"])
        for item in section["items"]:
            title, lines = item_to_lines(item)
            if title:
                out.append(f"**{title}**")
            for line in lines:
                out.append(f"- {line}")
    return "\n".join(out).rstrip() + "\n"


def render_html(resume: dict[str, Any]) -> str:
    body = [f"<h1>{html.escape(resume['name'])}</h1>"]
    if resume["headline"]:
        body.append(f"<p class=\"headline\">{html.escape(resume['headline'])}</p>")
    if resume["contact"]:
        body.append(f"<p class=\"contact\">{' · '.join(html.escape(x) for x in resume['contact'])}</p>")
    for section in resume["sections"]:
        body.append(f"<section><h2>{html.escape(section['title'])}</h2>")
        for item in section["items"]:
            title, lines = item_to_lines(item)
            if title:
                body.append(f"<h3>{html.escape(title)}</h3>")
            if lines:
                body.append("<ul>")
                body.extend(f"<li>{html.escape(line)}</li>" for line in lines)
                body.append("</ul>")
        body.append("</section>")
    return HTML_TEMPLATE.format(body="\n".join(body))


HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Resume</title>
<style>
  @page {{ size: A4; margin: 18mm; }}
  body {{
    color: #1f2933;
    font-family: Arial, "Noto Sans", "Microsoft YaHei", sans-serif;
    font-size: 10.5pt;
    line-height: 1.38;
    margin: 0 auto;
    max-width: 820px;
  }}
  h1 {{ font-size: 25pt; margin: 0 0 4px; letter-spacing: 0; }}
  .headline {{ color: #334e68; font-size: 11.5pt; margin: 0 0 4px; }}
  .contact {{ color: #52606d; font-size: 9.5pt; margin: 0 0 14px; }}
  h2 {{
    border-bottom: 1px solid #bcccdc;
    color: #102a43;
    font-size: 12pt;
    margin: 14px 0 7px;
    padding-bottom: 3px;
    text-transform: uppercase;
  }}
  h3 {{ font-size: 10.5pt; margin: 8px 0 3px; }}
  ul {{ margin: 3px 0 7px 18px; padding: 0; }}
  li {{ margin: 2px 0; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""


def write_docx(resume: dict[str, Any], path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise SystemExit("python-docx is required to write DOCX output.") from exc

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)

    document.add_heading(resume["name"], level=0)
    if resume["headline"]:
        document.add_paragraph(resume["headline"])
    if resume["contact"]:
        document.add_paragraph(" | ".join(resume["contact"]))

    for section_data in resume["sections"]:
        document.add_heading(section_data["title"], level=1)
        for item in section_data["items"]:
            title, lines = item_to_lines(item)
            if title:
                run = document.add_paragraph().add_run(title)
                run.bold = True
            for line in lines:
                document.add_paragraph(line, style="List Bullet")
    document.save(path)


def wrap_pdf_lines(text: str, max_chars: int) -> list[str]:
    text = str(text).strip()
    if not text:
        return [""]
    if re.search(r"\s", text):
        tokens = re.split(r"(\s+)", text)
        lines: list[str] = []
        current = ""
        for token in tokens:
            if len(current + token) > max_chars and current:
                lines.append(current.strip())
                current = token
            else:
                current += token
        if current.strip():
            lines.append(current.strip())
        return lines or [""]
    return [text[index : index + max_chars] for index in range(0, len(text), max_chars)]


def write_pdf(resume: dict[str, Any], path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise SystemExit("reportlab is required to write PDF output.") from exc

    width, height = A4
    margin = 52
    y = height - margin
    c = canvas.Canvas(str(path), pagesize=A4)

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        normal_font = "STSong-Light"
    except Exception:
        normal_font = "Helvetica"

    def new_page_if_needed(extra: float = 18) -> None:
        nonlocal y
        if y < margin + extra:
            c.showPage()
            y = height - margin

    def line_height(size: int) -> int:
        return size + 4

    def text_height(text: str, size: int = 10) -> float:
        max_chars = max(42, int((width - 2 * margin) / (size * 0.48)))
        return len(wrap_pdf_lines(text, max_chars)) * line_height(size)

    def draw_wrapped(text: str, size: int = 10, bullet: bool = False, bold: bool = False) -> None:
        nonlocal y
        font = "Helvetica-Bold" if bold and normal_font == "Helvetica" else normal_font
        c.setFont(font, size)
        prefix = "• " if bullet else ""
        max_chars = max(42, int((width - 2 * margin) / (size * 0.48)))
        lines = wrap_pdf_lines(text, max_chars)
        for index, line in enumerate(lines):
            new_page_if_needed(line_height(size) + 2)
            c.drawString(margin, y, (prefix if index == 0 else "  ") + line)
            y -= line_height(size)

    def block_height(title: str | None, lines: list[str]) -> float:
        height_needed = 2
        if title:
            height_needed += text_height(title, size=10)
        for line in lines:
            height_needed += text_height(line, size=10)
        return height_needed + 2

    c.setFont(normal_font, 24)
    c.drawString(margin, y, resume["name"])
    y -= 28
    if resume["headline"]:
        draw_wrapped(resume["headline"], size=11)
    if resume["contact"]:
        draw_wrapped(" | ".join(resume["contact"]), size=9)
    y -= 5

    for section_data in resume["sections"]:
        new_page_if_needed(36)
        c.setFont(normal_font, 12)
        c.drawString(margin, y, section_data["title"].upper())
        y -= 5
        c.line(margin, y, width - margin, y)
        y -= 14
        for item in section_data["items"]:
            title, lines = item_to_lines(item)
            needed = min(block_height(title, lines), height - (2 * margin))
            new_page_if_needed(needed)
            if title:
                draw_wrapped(title, size=10, bold=True)
            for line in lines:
                draw_wrapped(line, size=10, bullet=True)
            y -= 2
    c.save()


def unique_stem(output_dir: Path, base_name: str, suffixes: list[str]) -> str:
    candidate = base_name
    index = 1
    while any((output_dir / f"{candidate}{suffix}").exists() for suffix in suffixes):
        candidate = f"{base_name}_{index}"
        index += 1
    return candidate


def write_outputs(resume: dict[str, Any], output_dir: Path, base_name: str) -> str:
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = unique_stem(output_dir, base_name, [".md", ".html", ".docx", ".pdf"])
    (output_dir / f"{stem}.md").write_text(render_markdown(resume), encoding="utf-8")
    (output_dir / f"{stem}.html").write_text(render_html(resume), encoding="utf-8")
    write_docx(resume, output_dir / f"{stem}.docx")
    write_pdf(resume, output_dir / f"{stem}.pdf")
    return stem


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Export a resume to MD, HTML, DOCX, and PDF.")
    parser.add_argument("input", type=Path, help="Resume input file: JSON, YAML, or Markdown")
    parser.add_argument("--output-dir", type=Path, default=Path("resume-output"))
    parser.add_argument("--base-name", default="resume")
    args = parser.parse_args(argv)

    resume = load_input(args.input)
    stem = write_outputs(resume, args.output_dir, args.base_name)
    print(f"Wrote {stem}.md, .html, .docx, and .pdf to {args.output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
